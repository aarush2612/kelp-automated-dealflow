from docx import Document
from pathlib import Path
from typing import Dict, List


class CitationBuilder:
    """
    Builds a citation document mapping slides to source sections.
    """

    def __init__(self, company_slug: str, source_md_path: str):
        self.company_slug = company_slug
        self.source_md_path = source_md_path

    def build(self, canonical_data: Dict, insights: Dict, output_path: str):
        doc = Document()
        doc.add_heading(
            f"Citation Document – {self.company_slug} (Anonymized)", level=1
        )

        # Slide 1 citations
        self._slide_1(doc, canonical_data, insights)

        # Slide 2 citations
        self._slide_2(doc, canonical_data)

        # Slide 3 citations
        self._slide_3(doc, canonical_data)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    # ---------------- slide mappings ----------------

    def _slide_1(self, doc, data, insights):
        doc.add_heading("Slide 1 – Company Overview", level=2)

        self._add_claim(
            doc,
            "Business description and products",
            f"{self.source_md_path} → Business Description, Product & Services"
        )

        if data["company_profile"]["website"]:
            self._add_claim(
                doc,
                "Company website",
                data["company_profile"]["website"]
            )

        for insight in insights.get("highlights", []):
            self._add_claim(
                doc,
                insight,
                f"{self.source_md_path} → Key Operational Indicators"
            )

    def _slide_2(self, doc, data):
        doc.add_heading("Slide 2 – Financial Performance", level=2)

        self._add_claim(
            doc,
            "Revenue and PAT trends",
            f"{self.source_md_path} → Financials Status → Income Statement"
        )

    def _slide_3(self, doc, data):
        doc.add_heading("Slide 3 – SWOT Summary", level=2)

        self._add_claim(
            doc,
            "Strengths, Weaknesses, Opportunities, Threats",
            f"{self.source_md_path} → SWOT"
        )

    # ---------------- helpers ----------------

    def _add_claim(self, doc, claim: str, source: str):
        p = doc.add_paragraph()
        p.add_run("Claim: ").bold = True
        p.add_run(claim + "\n")
        p.add_run("Source: ").bold = True
        p.add_run(source)
